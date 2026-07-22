import base64
import io
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from PIL import Image

from app.core.errors import PublicError
from app.schemas.resume import ResumeDocument


TEMPLATES = {"ats", "campus", "experienced"}


def export_resume_docx(
    resume: ResumeDocument,
    template: str,
    photo_data_url: str | None = None,
) -> bytes:
    if template not in TEMPLATES:
        raise PublicError(400, "invalid_template", "请选择有效的简历模板")

    document = Document()
    _configure_document(document, template)
    photo = _decode_photo(photo_data_url) if template == "campus" else None
    _add_header(document, resume, template, photo)

    if template == "campus":
        order = ["education", "experience", "projects", "campus", "skills", "certifications", "summary"]
    elif template == "experienced":
        order = ["summary", "skills", "experience", "projects", "education", "certifications", "campus"]
    else:
        order = ["summary", "skills", "experience", "projects", "education", "campus", "certifications"]

    for section in order:
        _render_section(document, resume, section, template)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def safe_download_name(name: str) -> str:
    cleaned = re.sub(r"[^\w\u4e00-\u9fff-]+", "_", name, flags=re.UNICODE).strip("_")
    return (cleaned[:40] or "resume") + "_简历.docx"


def _configure_document(document: Document, template: str) -> None:
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(16)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(18 if template == "ats" else 17)
    section.right_margin = Mm(18 if template == "ats" else 17)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.08


def _add_header(document, resume, template, photo):
    contact = resume.contact
    if photo:
        table = document.add_table(rows=1, cols=2)
        table.autofit = False
        left, right = table.rows[0].cells
        left.width = Mm(145)
        right.width = Mm(28)
        name_paragraph = left.paragraphs[0]
        _add_name_runs(name_paragraph, contact.name, contact.target_role, template)
        image_run = right.paragraphs[0].add_run()
        image_run.add_picture(photo, width=Mm(25), height=Mm(35))
    else:
        name_paragraph = document.add_paragraph()
        _add_name_runs(name_paragraph, contact.name, contact.target_role, template)

    details = [contact.phone, contact.email, contact.city, *contact.links]
    details = [item for item in details if item]
    if details:
        paragraph = document.add_paragraph("  |  ".join(details))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(7)
        for run in paragraph.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(71, 85, 105)


def _add_name_runs(paragraph, name, target_role, template):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = paragraph.add_run(name or "姓名")
    name_run.bold = True
    name_run.font.size = Pt(20 if template == "ats" else 22)
    name_run.font.name = "Arial"
    _set_east_asia_font(name_run, "黑体")
    if target_role:
        role_run = paragraph.add_run(f"\n{target_role}")
        role_run.font.size = Pt(10)
        role_run.font.color.rgb = RGBColor(51, 65, 85)


def _render_section(document, resume, section, template):
    if section == "summary" and resume.summary:
        _add_section_heading(document, "个人概述", template)
        document.add_paragraph(resume.summary)
    elif section == "skills" and resume.skills:
        _add_section_heading(document, "核心技能", template)
        document.add_paragraph("  |  ".join(resume.skills))
    elif section == "education" and resume.education:
        _add_section_heading(document, "教育背景", template)
        for item in resume.education:
            title = " · ".join(value for value in [item.school, item.major, item.degree] if value)
            _add_entry(document, title, item.start_date, item.end_date, item.highlights)
    elif section == "experience" and resume.experience:
        _add_section_heading(document, "实习/工作经历", template)
        for item in resume.experience:
            title = " · ".join(value for value in [item.organization, item.role, item.location] if value)
            _add_entry(document, title, item.start_date, item.end_date, item.bullets)
    elif section == "projects" and resume.projects:
        _add_section_heading(document, "项目经历", template)
        for item in resume.projects:
            title = " · ".join(value for value in [item.name, item.role] if value)
            _add_entry(document, title, item.start_date, item.end_date, item.bullets)
    elif section == "campus" and resume.campus:
        _add_section_heading(document, "校园经历", template)
        _add_bullets(document, resume.campus)
    elif section == "certifications" and resume.certifications:
        _add_section_heading(document, "证书与奖项", template)
        _add_bullets(document, resume.certifications)


def _add_section_heading(document, title, template):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(12)
    _set_east_asia_font(run, "黑体")
    run.font.color.rgb = RGBColor(17, 24, 39) if template == "ats" else RGBColor(29, 78, 216)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "111827" if template == "ats" else "93C5FD")
    border.append(bottom)
    paragraph._p.get_or_add_pPr().append(border)


def _add_entry(document, title, start_date, end_date, bullets):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(3)
    title_run = paragraph.add_run(title or "经历")
    title_run.bold = True
    dates = " - ".join(value for value in [start_date, end_date] if value)
    if dates:
        date_run = paragraph.add_run(f"    {dates}")
        date_run.italic = True
        date_run.font.color.rgb = RGBColor(71, 85, 105)
    _add_bullets(document, bullets)


def _add_bullets(document, bullets):
    for bullet in bullets:
        if not bullet.strip():
            continue
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.left_indent = Mm(5)
        paragraph.paragraph_format.first_line_indent = Mm(-3)
        paragraph.add_run(bullet.strip())


def _decode_photo(data_url: str | None) -> io.BytesIO | None:
    if not data_url:
        return None
    match = re.fullmatch(r"data:image/(png|jpeg);base64,([A-Za-z0-9+/=\r\n]+)", data_url)
    if not match:
        raise PublicError(400, "invalid_photo", "照片数据格式无效")
    try:
        raw = base64.b64decode(match.group(2), validate=True)
        if len(raw) > 2 * 1024 * 1024:
            raise PublicError(413, "photo_too_large", "简历照片不能超过 2 MB")
        image = Image.open(io.BytesIO(raw))
        image.verify()
        image = Image.open(io.BytesIO(raw)).convert("RGB")
        output = io.BytesIO()
        image.save(output, format="PNG", optimize=True)
        output.seek(0)
        return output
    except PublicError:
        raise
    except Exception as exc:
        raise PublicError(400, "invalid_photo", "照片无法读取") from exc


def _set_east_asia_font(run, font_name: str) -> None:
    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.get_or_add_rFonts()
    fonts.set(qn("w:eastAsia"), font_name)
