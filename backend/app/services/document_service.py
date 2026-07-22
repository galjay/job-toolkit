import io
import zipfile
from pathlib import Path, PurePosixPath

import docx
import pdfplumber

from app.core.config import settings
from app.core.errors import PublicError


PDF_MIME_TYPES = {"application/pdf", "application/octet-stream"}
DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/zip",
    "application/octet-stream",
}


def parse_document(filename: str, content_type: str, content: bytes) -> dict:
    extension = Path(filename or "").suffix.lower()
    if extension not in {".pdf", ".docx"}:
        raise PublicError(400, "unsupported_file_type", "仅支持 PDF 和 DOCX 文件")
    if len(content) > settings.MAX_UPLOAD_BYTES:
        raise PublicError(413, "file_too_large", "文件不能超过 8 MB")

    if extension == ".pdf":
        if content_type not in PDF_MIME_TYPES:
            raise PublicError(400, "mime_type_mismatch", "文件类型与扩展名不一致")
        return _parse_pdf(content)

    if content_type not in DOCX_MIME_TYPES:
        raise PublicError(400, "mime_type_mismatch", "文件类型与扩展名不一致")
    return _parse_docx(content)


def _parse_pdf(content: bytes) -> dict:
    if not content.startswith(b"%PDF-"):
        raise PublicError(400, "invalid_file_signature", "PDF 文件签名无效")

    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            if len(pdf.pages) > settings.MAX_PDF_PAGES:
                raise PublicError(
                    400,
                    "too_many_pages",
                    f"PDF 不能超过 {settings.MAX_PDF_PAGES} 页",
                )
            text = "\n\n".join((page.extract_text() or "").strip() for page in pdf.pages)
            units = len(pdf.pages)
    except PublicError:
        raise
    except Exception as exc:
        raise PublicError(400, "corrupt_document", "PDF 无法读取或已经损坏") from exc

    return _document_result(text, "pdf", units)


def _parse_docx(content: bytes) -> dict:
    if not content.startswith(b"PK\x03\x04"):
        raise PublicError(400, "invalid_file_signature", "DOCX 文件签名无效")

    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            entries = archive.infolist()
            names = {entry.filename for entry in entries}
            if "[Content_Types].xml" not in names or "word/document.xml" not in names:
                raise PublicError(400, "invalid_file_signature", "文件不是有效的 DOCX")
            if any(_unsafe_archive_name(entry.filename) for entry in entries):
                raise PublicError(400, "unsafe_archive", "DOCX 包含不安全的归档路径")
            if any(entry.flag_bits & 0x1 for entry in entries):
                raise PublicError(400, "unsafe_archive", "不支持加密的 DOCX")
            if sum(entry.file_size for entry in entries) > settings.MAX_ARCHIVE_BYTES:
                raise PublicError(400, "unsafe_archive", "DOCX 解压后内容过大")

        document = docx.Document(io.BytesIO(content))
        chunks = [paragraph.text.strip() for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                chunks.append("\t".join(cell.text.strip() for cell in row.cells))
        text = "\n".join(chunk for chunk in chunks if chunk)
        units = len(document.paragraphs)
    except PublicError:
        raise
    except (zipfile.BadZipFile, ValueError, KeyError) as exc:
        raise PublicError(400, "corrupt_document", "DOCX 无法读取或已经损坏") from exc

    return _document_result(text, "docx", units)


def _unsafe_archive_name(name: str) -> bool:
    normalized = name.replace("\\", "/")
    path = PurePosixPath(normalized)
    return path.is_absolute() or ".." in path.parts


def _document_result(text: str, kind: str, units: int) -> dict:
    normalized = text.strip()
    if not normalized:
        raise PublicError(
            422,
            "no_extractable_text",
            "文件中没有可提取文本，扫描版 PDF 请先进行 OCR",
        )
    if len(normalized) > settings.MAX_DOCUMENT_CHARS:
        raise PublicError(
            400,
            "document_too_long",
            f"文档文字不能超过 {settings.MAX_DOCUMENT_CHARS} 字",
        )
    return {
        "text": normalized,
        "characters": len(normalized),
        "kind": kind,
        "units": units,
    }
