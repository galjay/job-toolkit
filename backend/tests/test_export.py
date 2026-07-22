import io

import docx
import pytest


RESUME = {
    "contact": {
        "name": "张三",
        "phone": "13800000000",
        "email": "zhangsan@example.com",
        "city": "天津",
        "target_role": "产品实习生",
        "links": ["github.com/example"],
    },
    "summary": "关注用户研究与数据分析，具备完整项目实践。",
    "education": [
        {
            "id": "edu-1",
            "school": "示例大学",
            "degree": "本科",
            "major": "环境工程",
            "start_date": "2022.09",
            "end_date": "2026.06",
            "highlights": ["主修课程：统计学、项目管理"],
        }
    ],
    "experience": [
        {
            "id": "exp-1",
            "organization": "示例科技",
            "role": "产品实习生",
            "location": "天津",
            "start_date": "2025.06",
            "end_date": "2025.09",
            "bullets": ["访谈 10 名用户并归纳需求，形成产品改进清单。"],
        }
    ],
    "projects": [
        {
            "id": "project-1",
            "name": "校园信息平台",
            "role": "项目负责人",
            "start_date": "2024.03",
            "end_date": "2024.06",
            "bullets": ["完成需求调研、原型设计与可用性测试。"],
        }
    ],
    "campus": ["学生会项目部成员"],
    "skills": ["用户研究", "Excel", "Axure"],
    "certifications": ["大学英语六级"],
}


def document_text(document):
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


@pytest.mark.parametrize("template", ["ats", "campus", "experienced"])
def test_docx_export_contains_editable_resume_text(client, template):
    response = client.post(
        "/api/resume/export/docx",
        json={"template": template, "resume": RESUME},
    )

    assert response.status_code == 200
    assert response.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    document = docx.Document(io.BytesIO(response.content))
    text = document_text(document)
    assert "张三" in text
    assert "项目经历" in text
    assert "校园信息平台" in text


def test_ats_export_uses_no_tables(client):
    response = client.post(
        "/api/resume/export/docx",
        json={"template": "ats", "resume": RESUME},
    )
    document = docx.Document(io.BytesIO(response.content))
    assert document.tables == []


def test_template_order_reflects_its_purpose(client):
    campus = docx.Document(
        io.BytesIO(
            client.post(
                "/api/resume/export/docx",
                json={"template": "campus", "resume": RESUME},
            ).content
        )
    )
    experienced = docx.Document(
        io.BytesIO(
            client.post(
                "/api/resume/export/docx",
                json={"template": "experienced", "resume": RESUME},
            ).content
        )
    )
    campus_text = document_text(campus)
    experienced_text = document_text(experienced)
    assert campus_text.index("教育背景") < campus_text.index("实习/工作经历")
    assert experienced_text.index("个人概述") < experienced_text.index("教育背景")


def test_rejects_unknown_template(client):
    response = client.post(
        "/api/resume/export/docx",
        json={"template": "decorative-neon", "resume": RESUME},
    )
    assert response.status_code == 400
    assert response.json()["code"] == "invalid_template"
